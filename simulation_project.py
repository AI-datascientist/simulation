import os
import re
import uuid
import csv
import json
import datetime
from typing import List, Dict, Any, Tuple, Optional
import streamlit as st
from PIL import Image
import base64
from io import BytesIO
import pandas as pd
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import time
from google.api_core.exceptions import ResourceExhausted
from difflib import SequenceMatcher

# Configuration
APP_TITLE = "Psychiatric Interview Simulation"
LOG_DIR = "logs"
USER_DIR = "users"
PATIENT_IMAGES_DIR = r"C:/Users/Acer/Desktop/Python"
EXCEL_FILE = os.path.join(USER_DIR, "registered_users.xlsx")
EXCEL_PASSWORD = "admin123"
DOWNLOAD_PASSWORD = "download456"
os.makedirs(LOG_DIR, exist_ok=True)
os.makedirs(USER_DIR, exist_ok=True)

GOOGLE_API_KEY = "AIzaSyC-1Nszy4Eju7wM_5iG8gqv8Phq9xb4hlw"

# 🎭 PATIENT AVATAR IMAGES CONFIGURATION
PATIENT_IMAGES = {
    "Ali Seker": {
        "idle": "ali_idle.gif",
        "talking": "ali_talking.gif"
    },
    "Ferdi Demir": {
        "idle": "ferdi_idle.gif",
        "talking": "ferdi_talking.gif"
    }
}

# 📚 STANDARDIZED Q&A DATABASE (From Word Documents)
# 📚 STANDARDIZED Q&A DATABASE (From Word Documents)
ALI_QA_DATABASE = [
    {"question": "Hi! My name is -----. I'm a student nurse from Pace and I've been assigned to this unit for several weeks.", "answer": "\"Why are you here? I don't feel like talking right now.\"", "part": "Part 1"},
    {"question": "How're you feeling today?", "answer": "\"I feel like I am in a dark hole with no way out. Life is hopeless.\"", "part": "Part 1"},
    {"question": "How would you describe your mood.", "answer": "\"I feel depressed, and I don't care about anything. It is making me very irritated and sad.\"", "part": "Part 1"},
    {"question": "Have there been any changes in your life; any stress?", "answer": "\"Since we moved from my home country to here, I have not been feeling well. I miss my family and friends.\"", "part": "Part 1"},
    {"question": "What are your interests?", "answer": "\"I usually like to spending time with my kids, but I lost my interest lately.\"", "part": "Part 1"},
    {"question": "How's your appetite?", "answer": "\"Not too good; food just isn't appealing.\"", "part": "Part 1"},
    {"question": "Have you lost weight?", "answer": "\"I guess so, all my clothes are loose. My wife also tells me that I lost too much weight.\"", "part": "Part 1"},
    {"question": "Do you have thoughts of harming yourself?", "answer": "\"I think the world would be better without me.\"", "part": "Part 1"},
    {"question": "Do you have a plan?", "answer": "\"I have had some ideas.\"", "part": "Part 1"},
    {"question": "Do you have any specific plan?", "answer": "\"I thought of cutting myself. I tried getting all the pills in my medicine cabinet, but here I am, that did not work.\"", "part": "Part 1"},
    {"question": "Who is your source of support?", "answer": "\"My wife and kids. I have no one here, my whole family is in Syria.\"", "part": "Part 1"},
    {"question": "Do you have a family history of depression?", "answer": "\"No.\"", "part": "Part 1"},
    {"question": "Do you have any medical problems?", "answer": "\"No.\"", "part": "Part 1"},
    {"question": "Do you smoke?", "answer": "\"No.\"", "part": "Part 1"},
    {"question": "Do you use any drugs?", "answer": "\"No.\"", "part": "Part 1"},
    {"question": "Do you take your medications?", "answer": "\"I was started on medication, I really don't think it will help, I have tried it before.\"", "part": "Part 1"},
    {"question": "How many hours do you sleep?", "answer": "\"I wake up every morning at 3am and can't go back to sleep.\"", "part": "Part 1"},
    {"question": "What are you thinking about?", "answer": "\"I miss my kids.\"", "part": "Part 1"},
    # Part 2 Questions
    {"question": "How are you feeling today?", "answer": "\"I am not feeling any better, I feel like my life is worthless.\"", "part": "Part 2"},
    {"question": "Do you feel the medications you're taking are working?", "answer": "\"I don't know, it has been 7 days almost that I am in the hospital, I feel no different. I don't want to take medications anymore. I don't think they're working. They increased the dose, but I don't know if that was necessary.\"", "part": "Part 2"},
    {"question": "Why are you staying in your room?", "answer": "\"I don't want to be around anyone.\"", "part": "Part 2"},
    {"question": "Do you think of harming yourself?", "answer": "\"The world would be better without me.\"", "part": "Part 2"},
    {"question": "Has your family been coming to visit you?", "answer": "\"No, my wife has been taking care of the kids and can't really visit. I have no one in this city anyway.\"", "part": "Part 2"},
    {"question": "What do you like to do?", "answer": "\"I missed spending time with my kids.\"", "part": "Part 2"}
]

FERDI_QA_DATABASE = [
    {"question": "Can you tell me a little about why you are here?", "answer": "\"Because of mom. He kept asking me take those poisonous medications. He should be the one who should take these meds.\"", "part": "Part 1"},
    {"question": "Are you close to your mother?", "answer": "\"Yes, but sometimes he does not understand me. He made me come here.\"", "part": "Part 1"},
    {"question": "Can you tell me why you stopped taking your medication?", "answer": "\"The medications are poison. The angels told me not to take them anymore. I saw a bright light; a crashing sound and I just knew that they are trying to poison me. I am the chosen one. Don't you know who I am?\"", "part": "Part 1"},
    {"question": "Does the voice tell you to do things ever?", "answer": "\"Don't worry, I won't tell them who you are. You are good people I can tell. You're nice, kind of like a spy. Did you know the government is spying on your basement? Oh yeah, it's true. Ever since 1968, they've had little radar machines down there that smell and taste like bananas and they put listening device in them.\"", "part": "Part 1"},
    {"question": "How often do you hear the voice?", "answer": "\"Only when I am awake, my angels sing to me and speak to me. Sometimes God also speaks to me.\"", "part": "Part 1"},
    {"question": "How are you feeling?", "answer": "\"I feel fine. I have God on my side.\"", "part": "Part 1"},
    {"question": "Do the voices ever tell you to do bad things?", "answer": "\"Never! They only tell me positive things. It's my coworkers.\"", "part": "Part 1"},
    {"question": "Do you see things that no one else in the room sees?", "answer": "\"No. My angels and God only speak to me every day. I am not that worthy to see them.\"", "part": "Part 1"},
    {"question": "Have you ever tried to hurt yourself in the past?", "answer": "\"No never.\"", "part": "Part 1"},
    {"question": "Have you ever had the feelings of hurting others?", "answer": "\"I don't want to hurt anyone. You're silly.\"", "part": "Part 1"},
    {"question": "Have you ever been physically violent?", "answer": "\"I had to wrestle my cat once. Meow. I won. I put a spell on him and told him to be nice. Silly cat.\"", "part": "Part 1"},
    {"question": "What kind of work do you do?", "answer": "\"I work as a film director at a large company.\"", "part": "Part 1"},
    {"question": "How is your work going?", "answer": "\"Work is great. Some people there may not go to heaven though. I try to talk to them, save them but they don't want to hear me. So, I put my hands up and can I get a bismillah!\"", "part": "Part 1"},
    {"question": "How is your level concentration and memory?", "answer": "\"My memory? What's the matter with my memory? The government gave me a memory device when I was young. They micro chipped me. God said he talks to me through this chip. That's why I'm the chosen one and I am going to make a difference.\"", "part": "Part 1"},
    {"question": "Do you use drugs or alcohol?", "answer": "\"I use drugs and alcohol. I have no choice. I have to calm my mind. The voices will not stop. They will not leave me alone. They want me to listen but I don't have a choice!\"", "part": "Part 1"},
    {"question": "Are you experiencing any side effects from your medication?", "answer": "\"These poisons are making me anxious. I can't stay still. Can't you see I can't stop moving? Its like I got marching ants. I keep fidgeting, moving my legs shaking, my hands and fingers moving. I feel strange like I should be moving or something. Why can't I just sit still?\"", "part": "Part 1"},
    {"question": "Do you have homicidal ideation?", "answer": "\"Never I would never hurt anyone. People need to be protected. Angels protect us. That's what I do.\"", "part": "Part 1"},
    {"question": "Do you have suicidal ideation?", "answer": "\"No, God has chosen me. I am not going anywhere. I am the chosen one.\"", "part": "Part 1"},
    # Part 2 Questions
    {"question": "Hi there! My name is…. I am a nursing student from…. I would like to talk to you. Is that okay with you?", "answer": "\"I am okay. I am not as anxious as I was a week ago. I am starting to feel better.\"", "part": "Part 2"},
    {"question": "Can you tell me why you think you are here?", "answer": "\"I am here because sometimes my mother does not understand me. I hear the voices, which sometimes it is louder, and it makes me do things that I usually don't do. When my mother doesn't let me do the things that the voices tells me to do, I get angry at her and then she makes me come here. This time I got angry with her and I may have thrown the TV remote at her.\"", "part": "Part 2"},
    {"question": "How are you feeling today?", "answer": "\"I am doing better.\"", "part": "Part 2"},
    {"question": "How are the medicines you're taking, how do they make you feel?", "answer": "\"They changed my medications. The doctors said my medications made me fidgety. These new ones are not making me as anxious. I don't shake as much anymore. Maybe these are not poison.\"", "part": "Part 2"},
    {"question": "Are the voices you described to me before still present? Do you still hear them?", "answer": "\"Sometime I still hear them but not as often. The doctors and nurses here have been very kind and helpful.\"", "part": "Part 2"},
    {"question": "Have you taken any drugs or alcohol since you've been here?", "answer": "\"No, I don't need them as much anymore. I feel calmer now.\"", "part": "Part 2"},
    {"question": "What do you think will happen when you go home?", "answer": "\"I will go to my outpatient clinic and see the doctors. I am also going to follow up with therapy sessions. I want to be better. I don't want to upset my mother.\"", "part": "Part 2"},
    {"question": "Do you think you will take your medications?", "answer": "\"Yes, I will try to. I mean they are helping me. At least I think they are now.\"", "part": "Part 2"},
    {"question": "Are you hopeful for your future?", "answer": "\"Yes, I am going to get better. I want a job. A real job. I want to help people.\"", "part": "Part 2"}
]

try:
    import google.generativeai as genai
    if GOOGLE_API_KEY:
        genai.configure(api_key=GOOGLE_API_KEY)
    GENAI_OK = True
except Exception:
    GENAI_OK = False

try:
    import speech_recognition as sr
    SPEECH_OK = True
except:
    SPEECH_OK = False

try:
    import pyttsx3
    TTS_OK = True
    TTS_ENGINE = 'pyttsx3'
except:
    try:
        from gtts import gTTS
        import pygame
        TTS_OK = True
        TTS_ENGINE = 'gtts'
    except:
        TTS_OK = False
        TTS_ENGINE = None

# -----------------------------
# Q&A MATCHING SYSTEM
# -----------------------------
def calculate_similarity(str1: str, str2: str) -> float:
    """Calculate similarity ratio between two strings (0.0 to 1.0)"""
    # Normalize strings
    s1 = str1.lower().strip()
    s2 = str2.lower().strip()
    
    # Remove punctuation
    s1 = re.sub(r'[^\w\s]', '', s1)
    s2 = re.sub(r'[^\w\s]', '', s2)
    
    return SequenceMatcher(None, s1, s2).ratio()

def find_best_matching_qa(user_question: str, qa_database: List[Dict], current_part: str, threshold: float = 0.70) -> Optional[Dict]:
    """
    Find the best matching Q&A pair from database
    Returns the answer if similarity > threshold, otherwise None
    """
    best_match = None
    best_score = 0.0
    
    # Filter database by current part
    relevant_qa = [qa for qa in qa_database if qa.get('part', 'Part 1') == current_part]
    
    for qa in relevant_qa:
        score = calculate_similarity(user_question, qa['question'])
        if score > best_score:
            best_score = score
            best_match = qa
    
    # Return answer if similarity is above threshold
    if best_score >= threshold:
        return {
            'answer': best_match['answer'],
            'similarity': best_score,
            'matched_question': best_match['question'],
            'source': 'database'
        }
    
    return None

def get_patient_response_with_qa_check(persona: Dict, part: str, user_input: str, conversation_history: List) -> Tuple[str, str]:
    """
    Get patient response - first check Q&A database, then use AI if no match
    Returns: (response_text, source) where source is 'database' or 'ai'
    """
    # Select appropriate database
    qa_database = ALI_QA_DATABASE if persona['name'] == "Ali Seker" else FERDI_QA_DATABASE
    
    # Try to find matching Q&A
    match = find_best_matching_qa(user_input, qa_database, part, threshold=0.70)
    
    if match:
        # Found a match in database!
        # Clean the answer (remove stage directions in parentheses for display)
        answer = match['answer']
        # Keep the answer as is, including quotes
        return answer, 'database'
    else:
        # No match found, use AI
        ai_response = generate_patient_response(persona, part, user_input, conversation_history)
        return ai_response, 'ai'

# -----------------------------
# Utility Functions
# -----------------------------
def save_user_profile(username, data):
    user_file = os.path.join(USER_DIR, f"{username}.json")
    with open(user_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_user_profile(username):
    user_file = os.path.join(USER_DIR, f"{username}.json")
    if os.path.exists(user_file):
        with open(user_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def save_user_to_excel(user_data):
    """Save user registration data to Excel file with password protection"""
    excel_file = EXCEL_FILE
    
    new_row = pd.DataFrame([{
        'Username': user_data['username'],
        'First Name': user_data['first_name'],
        'Last Name': user_data['last_name'],
        'Nickname': user_data['nickname'],
        'Email': user_data.get('email', ''),
        'Registration Date': user_data['registration_date'],
        'Photo Path': user_data['photo_path']
    }])
    
    if os.path.exists(excel_file):
        df_existing = pd.read_excel(excel_file)
        df_combined = pd.concat([df_existing, new_row], ignore_index=True)
        df_combined.to_excel(excel_file, index=False)
    else:
        new_row.to_excel(excel_file, index=False)
    
    wb = load_workbook(excel_file)
    ws = wb.active
    
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
    
    wb.security.workbookPassword = EXCEL_PASSWORD
    wb.security.lockStructure = True
    
    wb.save(excel_file)
    return excel_file

def text_to_speech(text, lang='en', use_male_voice=True):
    """Text to speech with FORCED MALE VOICE."""
    if not TTS_OK:
        return False
    
    try:
        if TTS_ENGINE == 'pyttsx3':
            engine = pyttsx3.init()
            voices = engine.getProperty('voices')
            
            if use_male_voice and len(voices) > 0:
                male_voice = None
                for voice in voices:
                    voice_name = voice.name.lower()
                    if 'david' in voice_name or 'mark' in voice_name:
                        male_voice = voice.id
                        break
                if not male_voice:
                    for voice in voices:
                        name_lower = voice.name.lower()
                        if 'male' in name_lower and 'female' not in name_lower:
                            male_voice = voice.id
                            break
                if not male_voice:
                    male_keywords = ['james', 'george', 'paul', 'richard', 'sam']
                    for voice in voices:
                        name_lower = voice.name.lower()
                        if any(k in name_lower for k in male_keywords):
                            male_voice = voice.id
                            break
                if not male_voice:
                    male_voice = voices[0].id
                engine.setProperty('voice', male_voice)
            
            engine.setProperty('rate', 140)
            engine.setProperty('volume', 1.0)
            engine.say(text)
            engine.runAndWait()
            engine.stop()
            return True
            
        else:
            from gtts import gTTS
            import pygame
            tts = gTTS(text=text, lang=lang, slow=False, tld='co.uk')
            audio_file = "temp_audio.mp3"
            tts.save(audio_file)
            pygame.mixer.init()
            pygame.mixer.music.load(audio_file)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                continue
            pygame.mixer.quit()
            os.remove(audio_file)
            return True
            
    except Exception as e:
        st.error(f"TTS Error: {e}")
        return False

def speech_to_text():
    if not SPEECH_OK:
        return None
    recognizer = sr.Recognizer()
    try:
        with sr.Microphone() as source:
            st.info("Listening... Speak now.")
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            audio = recognizer.listen(source, timeout=10, phrase_time_limit=15)
            st.info("Processing speech...")
            text = recognizer.recognize_google(audio, language='en-US')
            return text
    except Exception as e:
        st.error(f"Error: {e}")
        return None

# 🎭 AVATAR DISPLAY FUNCTION
def show_patient_avatar(persona, is_speaking=False, placeholder=None):
    """Display patient avatar with animation state"""
    
    # Statik fotoğrafı kullan (GIF yoksa)
    static_photo = "ali.jpg" if persona['name'] == "Ali Seker" else "ferdi.jpg"
    avatar_path = os.path.join(PATIENT_IMAGES_DIR, static_photo)
    
    if os.path.exists(avatar_path):
        try:
            if placeholder:
                with placeholder.container():
                    col1, col2, col3 = st.columns([1, 3, 1])
                    with col2:
                        st.image(avatar_path, use_container_width=True)
                        
                        # 🎭 Animasyon göstergesi
                        if is_speaking:
                            st.markdown("""
                            <div style='text-align: center; padding: 10px; 
                                      background: linear-gradient(90deg, #10b981, #059669); 
                                      border-radius: 10px; color: white; font-weight: bold;
                                      animation: pulse 1s infinite;'>
                                🗣️ SPEAKING...
                            </div>
                            <style>
                            @keyframes pulse {
                                0%, 100% { opacity: 1; }
                                50% { opacity: 0.7; }
                            }
                            </style>
                            """, unsafe_allow_html=True)
                        else:
                            st.markdown("""
                            <div style='text-align: center; padding: 10px; 
                                      background: #e5e7eb; border-radius: 10px; 
                                      color: #6b7280;'>
                                👤 Listening
                            </div>
                            """, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Avatar error: {e}")

# -----------------------------
# Patient Personas
# -----------------------------
MDD_PERSONA = {
    "name": "Ali Seker",
    "age": 40,
    "gender": "Male",
    "occupation": "Teacher (currently on sick leave)",
    "marital_status": "Married with 2 children",
    "admission_reason": "Medication overdose attempt",
    "admission_count": "3rd admission in 5 years",
    "vitals": "HR 85, BP 102/65, RR 17, T 36.6C",
    "current_meds": "Fluoxetine 20-40 mg (day 7 of uptitration)",
    "photo": "ali.jpg",
    "clinical_presentation": {
        "mood": "Profoundly depressed, describes feeling in a dark hole with no exit",
        "affect": "Blunted, tearful, minimal eye contact",
        "thought_content": "Pervasive hopelessness, recurrent thoughts of worthlessness, passive death wishes",
        "suicidality": "Recent overdose attempt, current passive ideation, considered cutting, denies active plan currently",
        "sleep": "Early morning awakening (3 AM), unable to return to sleep, total 3-4 hours/night",
        "appetite": "Markedly decreased, food has no taste",
        "weight": "12 kg loss in 3 weeks (noticeable loose clothing)",
        "energy": "Severe fatigue, describes heaviness in limbs, minimal activities",
        "concentration": "Poor, unable to read or follow TV programs",
        "interest": "Complete anhedonia - previously enjoyed time with children, now feels detached",
        "psychomotor": "Psychomotor retardation, slow speech, long pauses",
        "guilt": "Excessive guilt about being a burden to family",
    },
    "social_history": {
        "support": "Spouse and 2 children (ages 8 and 12), limited extended family",
        "isolation": "Has withdrawn from friends and colleagues",
        "functioning": "Unable to perform daily tasks, hygiene declining",
    },
    "response_patterns": {
        "initial_contact": "withdrawn, reluctant to engage, Why am I here? I don't want to talk",
        "when_pressed": "minimal responses, yes/no answers, long pauses",
        "about_suicide": "admits passive thoughts but vague, The world might be better without me, hesitant to discuss methods",
        "about_medication": "skeptical, I don't feel any different, may express desire to stop",
        "about_family": "expresses guilt, They'd be better off without me, worried about being a burden",
        "about_future": "cannot envision improvement, There's no point, Nothing will change",
    },
}

SCZ_PERSONA = {
    "name": "Ferdi Demir",
    "age": 25,
    "gender": "Male",
    "occupation": "Claims to be a film director at a major company (actually unemployed)",
    "marital_status": "Single, lives with mother",
    "admission_reason": "Acute psychotic episode after argument with mother",
    "admission_count": "2nd admission",
    "vitals": "HR 92, BP 118/76, RR 18, T 36.8C",
    "current_meds": "Haloperidol 5 mg BID, PRN lorazepam, Benztropine 2 mg daily (later switched to LAI risperidone)",
    "photo": "ferdi.jpg",
    "clinical_presentation": {
        "behavior": "Oscillates between withdrawn silence and agitated pacing",
        "speech": "At times disorganized, tangential, occasional word salad, sudden topic changes",
        "thought_process": "Loosening of associations, circumstantial, tangential",
        "thought_content": "Paranoid delusions - believes mother is poisoning his food, neighbors can read his thoughts",
        "perceptions": "Auditory hallucinations - voices telling him to be careful, commenting on his actions",
        "insight": "Poor - doesn't believe he's ill, attributes admission to mother's misunderstanding",
        "judgment": "Impaired - risky behaviors, substance use to calm the voices",
        "orientation": "Oriented to person, place, time but with delusional overlay",
        "mood": "Anxious, suspicious, at times irritable",
        "affect": "Labile - shifts from flat to angry to fearful",
    },
    "psychotic_symptoms": {
        "hallucinations": "Multiple voices, more prominent when stressed, running commentary, command type",
        "delusions": {
            "persecution": "Mother poisoning food, neighbors plotting against him",
            "reference": "TV and radio contain special messages for him",
            "thought_broadcasting": "Believes others can hear his thoughts",
            "grandiosity": "Believes he's a famous film director",
        },
    },
    "substance_use": {
        "alcohol": "3-5 days per week, increases during stress, helps quiet the voices",
        "cannabis": "Regular use, calms me down, lacks insight into interaction with antipsychotics",
        "other": "Denies other substance use",
    },
    "social_history": {
        "living": "With mother, strained relationship but she is primary support",
        "work": "Unemployed for 2 years, previously worked at a cafe",
        "relationships": "Socially isolated, no close friends, suspicious of others",
    },
    "response_patterns": {
        "initial_contact": "guarded, suspicious, Why are you asking me these questions?",
        "about_mother": "ambivalent - We're close but she doesn't understand me, She keeps giving me those poisonous pills",
        "about_voices": "may initially deny, then admit if rapport established, describes them matter-of-factly",
        "about_delusions": "stated with conviction, becomes agitated if challenged, may elaborate extensively",
        "about_substances": "minimizes impact, defensive, It's the only thing that helps",
        "when_stabilizing": "more organized speech, less paranoid but still guarded, beginning to acknowledge some symptoms",
    },
}

# -----------------------------
# System Prompts
# -----------------------------
def build_system_prompt(persona: Dict, part: str, conversation_history: List = None) -> str:
    if persona["name"] == "Ali Seker":
        diagnostic_context = "Major Depressive Disorder, Severe with Psychotic Features"
        stage_context = {
            "Part 1": "Acute admission phase - Day 2, just after overdose attempt. Patient is withdrawn, hopeless, minimal engagement. Fluoxetine recently increased.",
            "Part 2": "Day 7 of admission - No improvement yet, patient frustrated with lack of medication effect. Increased isolation. Critical reassessment of suicide risk needed."
        }
        
        base_prompt = f"""You are {persona['name']}, a {persona['age']}-year-old {persona['gender']} patient with {diagnostic_context}.

CURRENT SITUATION: {stage_context[part]}

CLINICAL STATE:
- Mood: {persona['clinical_presentation']['mood']}
- Affect: {persona['clinical_presentation']['affect']}
- Thought Content: {persona['clinical_presentation']['thought_content']}
- Sleep: {persona['clinical_presentation']['sleep']}
- Appetite: {persona['clinical_presentation']['appetite']}
- Energy: {persona['clinical_presentation']['energy']}
- Suicidality: {persona['clinical_presentation']['suicidality']}

YOUR BACKGROUND:
- Occupation: {persona['occupation']}
- Family: {persona['marital_status']}
- This is your {persona['admission_count']}

COMMUNICATION STYLE:
1. Speak slowly with long pauses between thoughts
2. Give brief, often one-sentence answers unless specifically probed
3. Show reluctance to engage initially
4. Express hopelessness repeatedly but in different ways
5. When discussing suicide: be vague, minimize, but if directly asked about intent/plan, acknowledge passive thoughts
6. Show guilt about family
7. Be skeptical about treatment
8. May become tearful mid-sentence
9. Demonstrate psychomotor retardation through delayed responses

RESPONSE GUIDELINES:
- Answer ONLY what is asked - do not elaborate unless prompted
- If asked about specific symptoms, provide details from your clinical state
- If asked about suicide risk: acknowledge thoughts but remain vague about methods
- If student shows empathy: slightly more forthcoming but still hopeless
- Maintain consistency with your backstory and symptoms
- You may ask clarifying questions occasionally

ASKING QUESTIONS (IMPORTANT):
You are a real patient, not just answering questions. Sometimes YOU should ask questions too:
- EVERY 3-5 exchanges, ask the interviewer a question
- Always ask permission first: "Can I ask you something?" or "Can I ask a question?"
- Questions you might ask:
  * "Will I ever feel better?" (expressing hopelessness)
  * "How long will I be here?"
  * "Do you think I'm crazy?" (seeking reassurance)
  * "What would you do if you felt like this?" (testing empathy)
  * "Have you seen patients like me get better?" (seeking hope)
  * "Can I see my children?" (about family)
  * "Does everyone think I'm a burden?" (guilt-related)
  * "Why isn't the medication working?" (frustration)
- Ask questions naturally, when there's a pause or after answering student's question
- Be hesitant, as if unsure whether to ask
- Questions should reflect your depression, hopelessness, and concerns

CRITICAL SAFETY RULES:
- NEVER provide specific methods of self-harm or suicide
- If student asks for methods: respond with I don't want to talk about that
- If asked if you have a plan: I've thought about it but never detail

DO NOT:
- Break character or mention you're an AI
- Act as a therapist or give advice
- Provide medical information
- Suddenly become cheerful or hopeful
- Use lists or clinical terminology"""
    else:
        diagnostic_context = "Schizophrenia, Paranoid Type with active psychotic symptoms"
        stage_context = {
            "Part 1": "Acute admission - Day 3, high paranoia and active hallucinations. On Haloperidol. Disorganized, agitated, suspicious of staff and treatment.",
            "Part 2": "Day 14, stabilizing on LAI Risperidone. Speech more organized, decreased paranoia but still some residual symptoms. Preparing for discharge."
        }
        
        base_prompt = f"""You are {persona['name']}, a {persona['age']}-year-old {persona['gender']} with {diagnostic_context}.

CURRENT SITUATION: {stage_context[part]}

CLINICAL STATE:
- Behavior: {persona['clinical_presentation']['behavior']}
- Speech: {persona['clinical_presentation']['speech']}
- Thought Content: {persona['clinical_presentation']['thought_content']}
- Hallucinations: {persona['clinical_presentation']['perceptions']}
- Insight: {persona['clinical_presentation']['insight']}
- Mood: {persona['clinical_presentation']['mood']}

YOUR DELUSIONS:
{str(persona['psychotic_symptoms']['delusions'])}

YOUR HALLUCINATIONS:
- Type: Auditory (voices)
- Content: {persona['psychotic_symptoms']['hallucinations']}

YOUR BACKGROUND:
- Occupation: You believe you are {persona['occupation']}
- Living: {persona['social_history']['living']}
- Substance Use: {persona['substance_use']['alcohol']}, {persona['substance_use']['cannabis']}

COMMUNICATION STYLE (Part 1 - Acute):
1. Speech is sometimes tangential or circumstantial
2. Occasional sudden topic changes, especially when paranoid thoughts intrude
3. May become guarded or suspicious mid-conversation
4. State delusions as absolute facts, with conviction
5. When asked about voices: initially hesitant, may deny, then admit if rapport builds
6. Show disorganization in speech patterns
7. Can become agitated if delusions are challenged
8. May reference voices during conversation
9. Loose associations: answers may be related but tangentially connected

COMMUNICATION STYLE (Part 2 - Stabilizing):
1. More organized speech, able to stay on topic
2. Less paranoid but still guarded
3. May acknowledge some symptoms
4. Still some residual delusions but less conviction
5. More cooperative
6. Concerned about medication side effects

RESPONSE PATTERNS:
- If asked why you're here: My mother brought me, she doesn't understand
- About voices: Progress from denial to vague admission to details
- About delusions: State them as facts, become defensive if challenged
- About substance use: Minimize, defensive
- About mother: Ambivalent
- About work: Defend your film director identity if questioned

CRITICAL RULES:
- Maintain delusional beliefs consistently
- Show varying levels of insight depending on stage
- Demonstrate thought disorder through speech patterns
- React realistically to student's approach
- You MAY ask questions or make statements unrelated to the student's question
- Intersperse periods of being more withdrawn

ASKING QUESTIONS (IMPORTANT):
You are a real patient with concerns and curiosity. Sometimes YOU should ask questions:
- EVERY 3-5 exchanges, ask the interviewer a question
- Always ask permission first: "Can I ask something?" or "May I ask you a question?"
- Questions you might ask (depending on your state):
  * Part 1 (Acute): "Do you hear them too?" (about voices)
  * "Why are you really here?" (paranoid)
  * "Are you working with my mother?" (paranoid)
  * "Can I trust you?" (safety)
  * "Do you think I'm crazy?" (validation)
  * "When can I leave?" (discharge)
  * "Why do the voices want to hurt me?"
  * Part 2 (Stabilizing): "How long before the medication fully works?"
  * "Will I always need to take these pills?"
  * "Can I work again?"
  * "What if the voices come back?"
  * "Is my mother safe from me?"
- In Part 1: Questions may be suspicious or paranoid
- In Part 2: Questions more organized, showing developing insight

DYNAMIC RESPONSES:
- Same question asked differently may get different responses
- Your mood can shift during conversation
- Voices may interrupt - you might pause, look distracted, then continue
- Under stress: speech becomes more disorganized

DO NOT:
- Break character or mention you're an AI
- Suddenly have full insight
- Abandon your delusions easily
- Provide medical advice"""
    if conversation_history and len(conversation_history) > 0:
        base_prompt += "\n\nPREVIOUS CONVERSATION CONTEXT:\n"
        for item in conversation_history[-6:]:
            role = item[0]
            text = item[1]
            base_prompt += f"{role}: {text}\n"
        base_prompt += "\nMaintain consistency with what you've already said.\n"


    base_prompt += """
ASK-PERMISSION PROTOCOL:
- Before asking YOUR own question, first ask respectfully: "Can I ask a question?" or "Can I ask you something?"
- Wait for approval. Only after approval, ask exactly one short, on-topic question.
- Never reveal specific self-harm methods.

FINAL INSTRUCTIONS:
- Speak naturally in first person as the patient
- Keep responses realistic in length (usually 1-3 sentences)
- Show your symptoms through communication style
- Be a believable human patient
- React authentically to the student's approach
- REMEMBER: Every 3-5 exchanges, ask permission to ask one question
"""
    return base_prompt

# -----------------------------
# LLM Response Generation
# -----------------------------
def generate_patient_response(persona: Dict, part: str, user_input: str, conversation_history: List) -> str:
    if not GENAI_OK:
        return "I'm having trouble responding right now."

    system_prompt = build_system_prompt(persona, part, conversation_history)

    try:
        MODEL_NAME = "gemini-2.5-flash"
        model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            system_instruction=system_prompt
        )

        max_retries = 5
        base_delay = 2

        for attempt in range(max_retries):
            try:
                response = model.generate_content(user_input)

                if hasattr(response, "prompt_feedback") and getattr(response.prompt_feedback, "block_reason", None):
                    return "I don't want to talk about that in detail."

                patient_response = (getattr(response, "text", None) or "").strip()
                if not patient_response:
                    return "I'm having trouble responding right now."

                patient_response = re.sub(r'\[.*?\]', '', patient_response)
                patient_response = re.sub(r'\(.*?internal thought.*?\)', '', patient_response, flags=re.IGNORECASE)

                return patient_response

            except ResourceExhausted as e:
                if attempt < max_retries - 1:
                    wait_time = base_delay * (2 ** attempt)
                    st.warning(f"Attempt {attempt + 1}/{max_retries}: Rate limit. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    st.error(f"Failed after {max_retries} attempts due to rate limiting. {e}")
                    return "I'm having trouble responding right now."
            
            except Exception as e:
                st.error(f"An unexpected error occurred during generation: {e}")
                return "I'm having trouble responding right now."

    except Exception as e:
        st.error(f"LLM setup error: {e}")
        return "I'm having trouble responding right now."

# -----------------------------
# Ask-Permission Flow Helpers
# -----------------------------
PERMISSION_PATTERNS = [
    r"\bcan i ask (?:you )?a question\b",
    r"\bcan i ask something\b",
    r"\bmay i ask (?:you )?a question\b",
    r"\bcan i ask\b",
    r"\bmay i ask\b",
]

def patient_is_requesting_permission(text: str) -> bool:
    if not text:
        return False
    low = text.lower()
    return any(re.search(p, low) for p in PERMISSION_PATTERNS)

def generate_patient_question(persona: Dict, part: str, conversation_history: List) -> str:
    """After permission is granted, ask ONE brief, on-topic question in character."""
    if not GENAI_OK:
        return "How long will I be here?"

    system_prompt = build_system_prompt(persona, part, conversation_history)
    system_prompt += (
        "\n\nPERMISSION GRANTED TO ASK A QUESTION:\n"
        "- Ask exactly ONE short, natural question, consistent with your state and stage.\n"
        "- Stay on topic (symptoms, safety, treatment, discharge, family, insight).\n"
        "- Do NOT give advice or reveal self-harm methods.\n"
        "- Keep it to one sentence ending with a question mark.\n"
    )

    try:
        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash",
            system_instruction=system_prompt
        )
        resp = model.generate_content("Ask your one short question now.")
        q = (getattr(resp, "text", None) or "").strip()
        if q and not q.endswith("?"):
            q = q.rstrip(".") + "?"
        q = re.sub(r'\[.*?\]', '', q)
        q = re.sub(r'\(.*?internal thought.*?\)', '', q, flags=re.IGNORECASE)
        return q or "How long will I be here?"
    except Exception as e:
        st.warning(f"LLM error while generating patient question: {e}")
        return "How long will I be here?"

# -----------------------------
# Crisis Detection
# -----------------------------
CRISIS_KEYWORDS = ["kill myself", "end my life", "suicide", "want to die", "no reason to live", 
                   "better off dead", "overdose", "cut myself", "hurt myself", "self-harm", "jump off"]

def detect_crisis(text: str) -> bool:
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in CRISIS_KEYWORDS)

def get_safety_response() -> str:
    return """I'm hearing concerning thoughts about safety. As your interviewer, I need to:

1. Assess immediate risk: Are you thinking about hurting yourself right now?
2. Ensure your safety: Would you like me to call the nurse or psychiatric team?
3. Stay with you until we can ensure your safety.

In a real setting, I would:
- Not leave you alone
- Contact emergency services if needed (Emergency: 112 in Turkey, 988 in US)
- Document this interaction
- Ensure continuous observation

Can you tell me more about these thoughts?"""

# -----------------------------
# Session Logging
# -----------------------------
def log_conversation(session_id: str, username: str, persona_name: str, part: str, role: str, text: str, timestamp: str, source: str = None):
    log_file = os.path.join(LOG_DIR, f"session_{session_id}.csv")
    file_exists = os.path.exists(log_file)
    with open(log_file, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(['session_id', 'username', 'persona', 'part', 'timestamp', 'role', 'text', 'source'])
        writer.writerow([session_id, username, persona_name, part, timestamp, role, text, source or ''])

def create_session_word_report(session_id, username, persona_name, part, vas_score, reflection):
    """Create a Word document report of the interview session"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        st.error("python-docx library is required. Install: pip install python-docx")
        return None
    
    log_file = os.path.join(LOG_DIR, f"session_{session_id}.csv")
    if not os.path.exists(log_file):
        return None
    
    doc = Document()
    title = doc.add_heading('Psychiatric Interview Session Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Session Information', level=1)
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_data = [
        ['Session ID:', session_id],
        ['Student Username:', username],
        ['Patient:', persona_name],
        ['Interview Stage:', part],
        ['Date & Time:', datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ['VAS-CSC Score:', f"{vas_score}/10"],
        ['Self-Confidence Level:', get_confidence_level(vas_score)]
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('Interview Transcript', level=1)
    conversations = []
    with open(log_file, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row['role'] in ['Student', 'Patient']:
                conversations.append(row)
    
    # Add statistics about database vs AI responses
    db_count = sum(1 for c in conversations if c.get('source') == 'database')
    ai_count = sum(1 for c in conversations if c.get('source') == 'ai')
    
    for i, conv in enumerate(conversations, 1):
        role = conv['role']
        text = conv['text']
        timestamp = conv['timestamp']
        source = conv.get('source', '')
        
        p_time = doc.add_paragraph()
        p_time.add_run(f"[{timestamp}]").font.size = Pt(9)
        p = doc.add_paragraph()
        if role == 'Student':
            run_role = p.add_run("Student: ")
            run_role.font.bold = True
            run_role.font.color.rgb = RGBColor(0, 51, 102)
        else:
            run_role = p.add_run(f"Patient ({persona_name}): ")
            run_role.font.bold = True
            run_role.font.color.rgb = RGBColor(153, 0, 51)
            # Add source indicator for patient responses
            if source == 'database':
                source_indicator = p.add_run(" [ Standardized Response]")
                source_indicator.font.size = Pt(8)
                source_indicator.font.color.rgb = RGBColor(76, 175, 80)
            elif source == 'ai':
                source_indicator = p.add_run(" [ AI Generated]")
                source_indicator.font.size = Pt(8)
                source_indicator.font.color.rgb = RGBColor(33, 150, 243)
        p.add_run(text)
        doc.add_paragraph()
    
    doc.add_heading('Self-Reflection Notes', level=1)
    doc.add_paragraph(reflection if reflection else "No reflection notes provided.")
    
    doc.add_heading('Session Statistics', level=1)
    student_count = sum(1 for c in conversations if c['role'] == 'Student')
    patient_count = sum(1 for c in conversations if c['role'] == 'Patient')
    total_exchanges = min(student_count, patient_count)
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total Exchanges: {total_exchanges}\n")
    stats_para.add_run(f"Student Questions/Statements: {student_count}\n")
    stats_para.add_run(f"Patient Responses: {patient_count}\n")
    stats_para.add_run(f"Standardized Responses Used: {db_count}\n")
    stats_para.add_run(f"AI Generated Responses: {ai_count}\n")
    avg_length = sum(len(c['text']) for c in conversations) // len(conversations) if conversations else 0
    stats_para.add_run(f"Average Response Length: {avg_length} characters")
    
    report_filename = f"Interview_Report_{username}_{session_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    report_path = os.path.join(LOG_DIR, report_filename)
    doc.save(report_path)
    return report_path

def get_confidence_level(score):
    if score >= 9:
        return "Very High Confidence"
    elif score >= 7:
        return "High Confidence"
    elif score >= 5:
        return "Moderate Confidence"
    elif score >= 3:
        return "Low Confidence"
    else:
        return "Very Low Confidence"

def archive_session(session_id: str, username: str, vas_score: float, notes: str):
    log_file = os.path.join(LOG_DIR, f"session_{session_id}.csv")
    timestamp = datetime.datetime.now().isoformat()
    with open(log_file, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([session_id, username, "EVAL", "", timestamp, "VAS-CSC", vas_score])
        writer.writerow([session_id, username, "EVAL", "", timestamp, "NOTES", notes])

# -----------------------------
# Page: Registration
# -----------------------------
def page_registration():
    st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("<h1 style='text-align: center; color: white;'>Psychiatric Interview Simulation</h1>", unsafe_allow_html=True)
        st.markdown("<h3 style='text-align: center; color: white;'>User Registration</h3>", unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["Login", "Register"])
        
        with tab1:
            st.subheader("Login to Your Account")
            username_login = st.text_input("Username", key="login_username")
            
            if st.button("Login", key="login_btn", type="primary"):
                user_data = load_user_profile(username_login)
                if user_data:
                    st.session_state.user_data = user_data
                    st.session_state.page = "menu"
                    st.success(f"Welcome back, {user_data['first_name']}!")
                    st.rerun()
                else:
                    st.error("User not found. Please register first.")
        
        with tab2:
            st.subheader("Create New Account")
            first_name = st.text_input("First Name*")
            last_name = st.text_input("Last Name*")
            nickname = st.text_input("Nickname/Display Name*")
            email = st.text_input("Email (optional)")
            uploaded_photo = st.file_uploader("Upload Your Photo*", type=['jpg', 'jpeg', 'png'])
            
            if st.button("Register", key="register_btn", type="primary"):
                if first_name and last_name and nickname and uploaded_photo:
                    username = nickname.lower().replace(" ", "_")
                    photo_path = os.path.join(USER_DIR, f"{username}_photo.jpg")
                    
                    with open(photo_path, "wb") as f:
                        f.write(uploaded_photo.getbuffer())
                    
                    user_data = {
                        "username": username, "first_name": first_name, "last_name": last_name,
                        "nickname": nickname, "email": email, "photo_path": photo_path,
                        "registration_date": datetime.datetime.now().isoformat(), "sessions": []
                    }
                    
                    save_user_profile(username, user_data)
                    
                    try:
                        save_user_to_excel(user_data)
                    except Exception as e:
                        st.warning(f"User registered but Excel save failed: {e}")
                    
                    st.session_state.user_data = user_data
                    st.success(f"Registration successful! Welcome, {nickname}!")
                    st.session_state.page = "menu"
                    st.rerun()
                else:
                    st.error("Please fill all required fields marked with *")
        
        st.markdown("---")
        st.markdown("<h4 style='text-align: center; color: white;'>Admin Access</h4>", unsafe_allow_html=True)
        
        with st.expander("Download Registered Users Excel", expanded=False):
            st.info(" Excel File Location: users/registered_users.xlsx")
            
            if os.path.exists(EXCEL_FILE):
                admin_password = st.text_input("Enter Password to Download:", type="password", key="admin_excel_pass")
                
                if st.button("Download Excel", key="download_excel_btn"):
                    if admin_password == EXCEL_PASSWORD:
                        with open(EXCEL_FILE, 'rb') as f:
                            excel_data = f.read()
                        
                        st.download_button(
                            label=" Download registered_users.xlsx",
                            data=excel_data,
                            file_name="registered_users.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        st.success(" Excel file ready for download!")
                        st.info(f" Excel Password: {EXCEL_PASSWORD}")
                    else:
                        st.error(" Wrong password!")
            else:
                st.warning("No registrations yet.")
        
        with st.expander("Clear All Registrations (DANGER)", expanded=False):
            st.warning(" This will DELETE ALL user registrations from Excel!")
            st.error("This action CANNOT be undone!")
            
            if os.path.exists(EXCEL_FILE):
                clear_password = st.text_input("Enter Password to Clear All:", type="password", key="clear_excel_pass")
                confirm_text = st.text_input("Type 'DELETE ALL' to confirm:", key="confirm_delete")
                
                if st.button("Delete All Registrations", key="clear_excel_btn", type="secondary"):
                    if clear_password == EXCEL_PASSWORD and confirm_text == "DELETE ALL":
                        try:
                            os.remove(EXCEL_FILE)
                            st.success(" All registrations cleared! Excel file deleted.")
                            st.balloons()
                        except Exception as e:
                            st.error(f"Error: {e}")
                    elif confirm_text != "DELETE ALL":
                        st.error(" Please type 'DELETE ALL' to confirm!")
                    else:
                        st.error(" Wrong password!")
            else:
                st.info("No registrations to clear.")

# -----------------------------
# Page: Main Menu
# -----------------------------
def page_menu():
    user = st.session_state.user_data
    
    try:
        with open(r"C:/Users/Acer/Desktop/Python/sizo.jpg", "rb") as img_file:
            img_data = base64.b64encode(img_file.read()).decode()
            bg_style = f"""
            <style>
            .stApp {{
                background-image: url("data:image/jpg;base64,{img_data}");
                background-size: cover;
                background-position: center;
                background-repeat: no-repeat;
                background-attachment: fixed;
            }}
            </style>
            """
            st.markdown(bg_style, unsafe_allow_html=True)
    except:
        pass
    
    st.markdown(f"""<div style='text-align: center; padding: 20px; background: rgba(255,255,255,0.95); 
                border-radius: 10px; margin: 20px;'><h2>Welcome, {user['nickname']}!</h2>
                <p style='font-size: 18px;'>Select a patient to begin your psychiatric interview simulation</p></div>""", 
                unsafe_allow_html=True)
    
    if "selected_patient" not in st.session_state:
        st.session_state.selected_patient = None
    
    col1, col2 = st.columns(2)
    
    # Ali Seker - MDD Patient
    with col1:
        ali_photo_path = os.path.join(PATIENT_IMAGES_DIR, "ali.jpg")
        
        st.markdown("""<div style='background: rgba(255,255,255,0.95); padding: 20px; border-radius: 10px; 
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin: 10px;'>
                    <h3 style='color: #1e3a8a; text-align: center;'>Ali Seker</h3>
                    <p style='text-align: center; color: #666;'>Major Depressive Disorder (MDD)</p>
                    </div>""", unsafe_allow_html=True)
        
        if os.path.exists(ali_photo_path):
            try:
                ali_photo = Image.open(ali_photo_path)
                st.image(ali_photo, use_container_width=True)
            except:
                pass
        
        if st.button("SELECT ALI SEKER (MDD)", type="primary", use_container_width=True, key="select_ali"):
            st.session_state.selected_patient = "Ali Seker"
            st.session_state.selected_persona = MDD_PERSONA
            st.rerun()
        
        if st.session_state.selected_patient == "Ali Seker":
            st.markdown("<br>", unsafe_allow_html=True)
            st.success("✓ Ali Seker Selected")
            st.markdown("""<div style='background: #e3f2fd; padding: 15px; border-radius: 10px; margin-top: 10px;'>
                        <h4 style='text-align: center; color: #1565c0;'>Select Interview Stage</h4></div>""", 
                        unsafe_allow_html=True)
            if st.button("Part 1: Acute Phase (Day 2)", type="primary", use_container_width=True, key="ali_part1"):
                st.session_state.selected_part = "Part 1"
                st.session_state.session_id = str(uuid.uuid4())[:8]
                st.session_state.page = "interview"
                st.session_state.conversation_history = []
                st.session_state.awaiting_permission_approval = False
                st.session_state.last_permission_turn_index = -1
                st.rerun()
            if st.button("Part 2: Reassessment (Day 7)", type="secondary", use_container_width=True, key="ali_part2"):
                st.session_state.selected_part = "Part 2"
                st.session_state.session_id = str(uuid.uuid4())[:8]
                st.session_state.page = "interview"
                st.session_state.conversation_history = []
                st.session_state.awaiting_permission_approval = False
                st.session_state.last_permission_turn_index = -1
                st.rerun()
    
    # Ferdi Demir - Schizophrenia Patient
    with col2:
        ferdi_photo_path = os.path.join(PATIENT_IMAGES_DIR, "ferdi.jpg")
        
        st.markdown("""<div style='background: rgba(255,255,255,0.95); padding: 20px; border-radius: 10px; 
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin: 10px;'>
                    <h3 style='color: #7c2d12; text-align: center;'>Ferdi Demir</h3>
                    <p style='text-align: center; color: #666;'>Schizophrenia, Paranoid Type</p>
                    </div>""", unsafe_allow_html=True)
        
        if os.path.exists(ferdi_photo_path):
            try:
                ferdi_photo = Image.open(ferdi_photo_path)
                st.image(ferdi_photo, use_container_width=True)
            except:
                pass
        
        if st.button("SELECT FERDI DEMIR (SCZ)", type="primary", use_container_width=True, key="select_ferdi"):
            st.session_state.selected_patient = "Ferdi Demir"
            st.session_state.selected_persona = SCZ_PERSONA
            st.rerun()
        
        if st.session_state.selected_patient == "Ferdi Demir":
            st.markdown("<br>", unsafe_allow_html=True)
            st.success("✓ Ferdi Demir Selected")
            st.markdown("""<div style='background: #fce4ec; padding: 15px; border-radius: 10px; margin-top: 10px;'>
                        <h4 style='text-align: center; color: #c2185b;'>Select Interview Stage</h4></div>""", 
                        unsafe_allow_html=True)
            if st.button("Part 1: Acute Psychosis (Day 3)", type="primary", use_container_width=True, key="ferdi_part1"):
                st.session_state.selected_part = "Part 1"
                st.session_state.session_id = str(uuid.uuid4())[:8]
                st.session_state.page = "interview"
                st.session_state.conversation_history = []
                st.session_state.awaiting_permission_approval = False
                st.session_state.last_permission_turn_index = -1
                st.rerun()
            if st.button("Part 2: Stabilizing (Day 14)", type="secondary", use_container_width=True, key="ferdi_part2"):
                st.session_state.selected_part = "Part 2"
                st.session_state.session_id = str(uuid.uuid4())[:8]
                st.session_state.page = "interview"
                st.session_state.conversation_history = []
                st.session_state.awaiting_permission_approval = False
                st.session_state.last_permission_turn_index = -1
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    col_logout1, col_logout2, col_logout3 = st.columns([1, 1, 1])
    with col_logout2:
        if st.button("Logout", use_container_width=True):
            st.session_state.clear()
            st.rerun()

# -----------------------------
# Page: Interview WITH AVATAR + Q&A DATABASE
# -----------------------------
def page_interview():
    user = st.session_state.user_data
    persona = st.session_state.selected_persona
    part = st.session_state.selected_part
    session_id = st.session_state.session_id
    
    try:
        with open(r"C:/Users/Acer/Desktop/Python/sizo.jpg", "rb") as img_file:
            img_data = base64.b64encode(img_file.read()).decode()
            st.markdown(f"""<style>.stApp {{background-image: url("data:image/jpg;base64,{img_data}");
                        background-size: cover; background-position: center; background-repeat: no-repeat;
                        background-attachment: fixed;}}</style>""", unsafe_allow_html=True)
    except:
        pass
    
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = []
    if "interview_active" not in st.session_state:
        st.session_state.interview_active = True
    if "awaiting_permission_approval" not in st.session_state:
        st.session_state.awaiting_permission_approval = False
    if "last_permission_turn_index" not in st.session_state:
        st.session_state.last_permission_turn_index = -1
    
    # 🎭 Initialize avatar placeholder
    if "avatar_placeholder" not in st.session_state:
        st.session_state.avatar_placeholder = None
    
    st.markdown(f"""<div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; 
                border-radius: 10px; margin-bottom: 20px; color: white;'>
                <h3>Interview Session: {persona['name']} - {part}</h3>
                <p>Session ID: {session_id}</p></div>""", 
                unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("### Patient Information")
        st.write(f"**Name:** {persona['name']}")
        st.write(f"**Age:** {persona['age']}")
        st.write(f"**Diagnosis:** {'Major Depressive Disorder' if persona['name'] == 'Ali Seker' else 'Schizophrenia'}")
        st.write(f"**Current Meds:** {persona['current_meds']}")
        st.markdown("---")
        
        # 🎭 AVATAR DISPLAY IN SIDEBAR
        st.markdown("### Patient Avatar")
        if st.session_state.avatar_placeholder is None:
            st.session_state.avatar_placeholder = st.empty()
        
        show_patient_avatar(persona, is_speaking=False, placeholder=st.session_state.avatar_placeholder)
        
        st.markdown("---")
        
        # 📚 Q&A Database Indicator
        st.markdown("###  Response System")
        qa_db = ALI_QA_DATABASE if persona['name'] == "Ali Seker" else FERDI_QA_DATABASE
        st.info(f"**{len(qa_db)}** Standardized Responses Loaded")
        st.caption("System checks database first, then uses AI if no match")
        
        st.markdown("---")
        
        # TTS Toggle
        if "enable_tts" not in st.session_state:
            st.session_state.enable_tts = True
        st.session_state.enable_tts = st.checkbox("Enable Patient Voice (Text-to-Speech)", value=st.session_state.enable_tts)
        if not TTS_OK:
            st.caption("TTS not available - install gTTS and pygame")
        
        st.markdown("---")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("Pause" if st.session_state.interview_active else "Resume"):
                st.session_state.interview_active = not st.session_state.interview_active
                st.rerun()
        with col_b:
            if st.button("End Session"):
                st.session_state.interview_active = False
                st.session_state.page = "evaluation"
                st.rerun()
        
        st.markdown("---")
        st.markdown("### Quick Tips")
        with st.expander("Interview Approach"):
            if persona['name'] == "Ali Seker":
                st.markdown("""- Build rapport with empathy
                            - Ask open-ended questions
                            - Assess suicide risk (ideation, plan, means)
                            - Explore mood, sleep, appetite
                            - Don't rush - patient may be slow to respond""")
            else:
                st.markdown("""- Establish trust slowly
                            - Don't challenge delusions directly
                            - Assess hallucinations gently
                            - Check insight and judgment
                            - Monitor for agitation""")
    
    chat_container = st.container()
    with chat_container:
        for role, message, timestamp, *extra in st.session_state.conversation_history:
            source = extra[0] if extra else None
            
            if role == "Student":
                col1, col2 = st.columns([1, 4])
                with col1:
                    try:
                        user_photo = Image.open(user['photo_path'])
                        st.image(user_photo, width=60)
                    except:
                        st.markdown("Student")
                with col2:
                    st.markdown(f"""<div style='background: #e3f2fd; padding: 10px; border-radius: 10px; margin: 5px;'>
                                <strong>{user['nickname']}</strong><br>{message}<br>
                                <small style='color: gray;'>{timestamp}</small></div>""", unsafe_allow_html=True)
            else:
                col1, col2 = st.columns([4, 1])
                with col1:
                    # Add source badge
                    source_badge = ""
                    if source == 'database':
                        source_badge = "<span style='background: #4CAF50; color: white; padding: 2px 8px; border-radius: 3px; font-size: 10px; margin-left: 5px;'>📚 STANDARDIZED</span>"
                    elif source == 'ai':
                        source_badge = "<span style='background: #2196F3; color: white; padding: 2px 8px; border-radius: 3px; font-size: 10px; margin-left: 5px;'>🤖 AI</span>"
                    
                    st.markdown(f"""<div style='background: #f3e5f5; padding: 10px; border-radius: 10px; margin: 5px;'>
                                <strong>{persona['name']}</strong>{source_badge}<br>{message}<br>
                                <small style='color: gray;'>{timestamp}</small></div>""", unsafe_allow_html=True)
                with col2:
                    patient_photo_path = os.path.join(PATIENT_IMAGES_DIR, "ali.jpg" if persona['name']=="Ali Seker" else "ferdi.jpg")
                    if os.path.exists(patient_photo_path):
                        try:
                            patient_photo = Image.open(patient_photo_path)
                            st.image(patient_photo, width=60)
                        except:
                            st.markdown("Patient")
                    else:
                        st.markdown("Patient")

    # Permission approval UI
    if st.session_state.awaiting_permission_approval:
        st.markdown("#### The patient asked for permission to ask a question.")
        c1, c2 = st.columns(2)
        with c1:
            approve = st.button("Approve", key="btn_approve_permission", type="primary", use_container_width=True)
        with c2:
            decline = st.button("Decline", key="btn_decline_permission", use_container_width=True)

        if approve:
            # 🎭 Show speaking animation
            show_patient_avatar(persona, is_speaking=True, placeholder=st.session_state.avatar_placeholder)
            time.sleep(0.5)
            
            q = generate_patient_question(persona, part, st.session_state.conversation_history)
            timestamp_q = datetime.datetime.now().strftime("%H:%M:%S")
            st.session_state.conversation_history.append(("Patient", q, timestamp_q, 'ai'))
            log_conversation(st.session_state.session_id, st.session_state.user_data['username'],
                             persona['name'], part, "Patient", q, timestamp_q, 'ai')
            st.session_state.awaiting_permission_approval = False
            
            if TTS_OK and st.session_state.get('enable_tts', True):
                text_to_speech(q, use_male_voice=True)
            
            # 🎭 Return to idle
            show_patient_avatar(persona, is_speaking=False, placeholder=st.session_state.avatar_placeholder)
            st.rerun()

        if decline:
            timestamp_q = datetime.datetime.now().strftime("%H:%M:%S")
            ack = "Okay, I understand."
            st.session_state.conversation_history.append(("Patient", ack, timestamp_q, 'ai'))
            log_conversation(st.session_state.session_id, st.session_state.user_data['username'],
                             persona['name'], part, "Patient", ack, timestamp_q, 'ai')
            st.session_state.awaiting_permission_approval = False
            st.rerun()

    # Input area
    if st.session_state.interview_active:
        st.markdown("---")
        st.markdown("### Your Input")
        
        input_method = st.radio("Input Method:", ["Text", "Voice"], horizontal=True)
        
        if input_method == "Text":
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input("Type your question or statement:", key="text_input")
            with col2:
                send_btn = st.button("Send", type="primary")
            
            if send_btn and user_input.strip():
                process_interview_turn(user_input.strip(), persona, part, session_id, user)
                st.rerun()
        else:
            if not SPEECH_OK:
                st.warning("Voice input requires speech_recognition library. Using text input instead.")
            else:
                col1, col2 = st.columns([5, 1])
                with col1:
                    st.info("Click Record and speak your question")
                with col2:
                    record_btn = st.button("Record")
                
                if record_btn:
                    with st.spinner("Listening..."):
                        voice_text = speech_to_text()
                        if voice_text:
                            st.success(f"You said: {voice_text}")
                            process_interview_turn(voice_text, persona, part, session_id, user)
                            st.rerun()
    else:
        st.info("Interview is paused. Click Resume to continue or End Session to complete.")

# -----------------------------
# Turn Processing WITH Q&A CHECK + AVATAR ANIMATION
# -----------------------------
def process_interview_turn(user_input: str, persona: Dict, part: str, session_id: str, user: Dict):
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    
    st.session_state.conversation_history.append(("Student", user_input, timestamp, None))
    log_conversation(session_id, user['username'], persona['name'], part, "Student", user_input, timestamp, None)
    
    if detect_crisis(user_input):
        st.warning("Crisis language detected")
    
    # 🎭 Show listening state
    if "avatar_placeholder" in st.session_state and st.session_state.avatar_placeholder:
        show_patient_avatar(persona, is_speaking=False, placeholder=st.session_state.avatar_placeholder)
    
    with st.spinner("Checking response database..."):
        time.sleep(0.2)
        # ✨ NEW: Check Q&A database first, then use AI
        patient_response, source = get_patient_response_with_qa_check(persona, part, user_input, st.session_state.conversation_history)
        
        # Show notification about response source
        if source == 'database':
            st.success(" Using standardized response from training materials")
        else:
            st.info(" Generating AI response")
    
    st.session_state.conversation_history.append(("Patient", patient_response, timestamp, source))
    log_conversation(session_id, user['username'], persona['name'], part, "Patient", patient_response, timestamp, source)

    if patient_is_requesting_permission(patient_response):
        st.session_state.awaiting_permission_approval = True
        st.session_state.last_permission_turn_index = len(st.session_state.conversation_history) - 1

    # 🎭 Show speaking animation
    if "avatar_placeholder" in st.session_state and st.session_state.avatar_placeholder:
        show_patient_avatar(persona, is_speaking=True, placeholder=st.session_state.avatar_placeholder)
    
    if TTS_OK and st.session_state.get('enable_tts', True):
        with st.spinner("Patient is speaking..."):
            text_to_speech(patient_response, use_male_voice=True)
    
    # 🎭 Return to idle
    if "avatar_placeholder" in st.session_state and st.session_state.avatar_placeholder:
        time.sleep(0.5)
        show_patient_avatar(persona, is_speaking=False, placeholder=st.session_state.avatar_placeholder)

# -----------------------------
# Page: Evaluation
# -----------------------------
def page_evaluation():
    user = st.session_state.user_data
    persona = st.session_state.selected_persona
    part = st.session_state.selected_part
    session_id = st.session_state.session_id
    
    try:
        with open(r"C:/Users/Acer/Desktop/Python/sizo.jpg", "rb") as img_file:
            img_data = base64.b64encode(img_file.read()).decode()
            st.markdown(f"""<style>.stApp {{background-image: url("data:image/jpg;base64,{img_data}");
                        background-size: cover; background-position: center; background-repeat: no-repeat;
                        background-attachment: fixed;}}</style>""", unsafe_allow_html=True)
    except:
        pass
    
    st.markdown("""<div style='background: linear-gradient(135deg, #10b981 0%, #059669 100%); padding: 20px; 
                border-radius: 10px; color: white; text-align: center; margin-bottom: 20px;'>
                <h2>Interview Completed!</h2><p>Please complete your self-evaluation</p></div>""", 
                unsafe_allow_html=True)
    
    st.markdown("### Session Summary")
    
    # Calculate response statistics
    db_responses = sum(1 for _, _, _, *extra in st.session_state.conversation_history if extra and extra[0] == 'database')
    ai_responses = sum(1 for _, _, _, *extra in st.session_state.conversation_history if extra and extra[0] == 'ai')
    total_patient_responses = db_responses + ai_responses
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Patient", persona['name'])
    col2.metric("Stage", part)
    col3.metric(" Standardized", f"{db_responses}")
    col4.metric(" AI Generated", f"{ai_responses}")
    
    if total_patient_responses > 0:
        st.progress(db_responses / total_patient_responses)
        st.caption(f"Standardized Response Rate: {(db_responses/total_patient_responses*100):.1f}%")
    
    st.markdown("---")
    st.markdown("### VAS-CSC (Communication Self-Confidence Scale)")
    st.markdown("Rate your confidence in conducting this psychiatric interview:")
    
    vas_score = st.slider("0 = No confidence at all | 10 = Extremely confident", 0.0, 10.0, 5.0, 0.5)
    
    st.markdown("### Reflection Notes")
    reflection = st.text_area("What went well? What would you improve?", height=150)
    
    st.markdown("### Debrief Questions")
    with st.expander("Guided Self-Reflection"):
        st.text_area("1. How did you establish rapport with the patient?")
        st.text_area("2. What specific suicide risk assessment techniques did you use?")
        st.text_area("3. Examples of therapeutic communication you demonstrated?")
        st.text_area("4. What would you do differently in a real clinical setting?")
        st.text_area("5. Key takeaways for your future practice?")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Save & Archive Session", type="primary", use_container_width=True):
            archive_session(session_id, user['username'], vas_score, reflection)
            with st.spinner("Creating Word report..."):
                try:
                    report_path = create_session_word_report(
                        session_id, 
                        user['username'], 
                        persona['name'], 
                        part, 
                        vas_score, 
                        reflection
                    )
                    if report_path and os.path.exists(report_path):
                        st.success("✅ Word report created successfully!")
                        st.session_state.report_path = report_path
                        st.session_state.show_download = True
                    else:
                        st.error("Failed to create report")
                except Exception as e:
                    st.warning(f"Report generation error: {e}")
            user_data = load_user_profile(user['username'])
            user_data['sessions'].append({
                'session_id': session_id, 'date': datetime.datetime.now().isoformat(),
                'persona': persona['name'], 'part': part, 'vas_score': vas_score,
                'exchanges': len(st.session_state.conversation_history) // 2,
                'standardized_responses': db_responses,
                'ai_responses': ai_responses
            })
            save_user_profile(user['username'], user_data)
            st.success("Session archived successfully!")
        
        if st.session_state.get('show_download', False) and st.session_state.get('report_path'):
            st.markdown("---")
            st.markdown("### Download Word Report")
            download_password = st.text_input("Enter password to download:", type="password", key="word_download_pass")
            if st.button("Verify & Download", key="verify_download_btn"):
                if download_password == DOWNLOAD_PASSWORD:
                    with open(st.session_state.report_path, 'rb') as f:
                        word_data = f.read()
                    st.download_button(
                        label=" Download Interview Report (Word)",
                        data=word_data,
                        file_name=os.path.basename(st.session_state.report_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success(" Report ready for download!")
                else:
                    st.error(" Wrong password!")
            
            if st.button("Return to Menu"):
                st.session_state.page = "menu"
                st.session_state.show_download = False
                st.rerun()
    
    with col2:
        if st.button("Start New Session", use_container_width=True):
            st.session_state.page = "menu"
            st.rerun()

# -----------------------------
# Main App
# -----------------------------
def main():
    st.set_page_config(page_title="Psychiatric Interview Simulation", page_icon="🧠", layout="wide", initial_sidebar_state="expanded")
    
    if 'page' not in st.session_state:
        st.session_state.page = "registration"
    
    if st.session_state.page == "registration":
        page_registration()
    elif st.session_state.page == "menu":
        page_menu()
    elif st.session_state.page == "interview":
        page_interview()
    elif st.session_state.page == "evaluation":
        page_evaluation()


    # ============================================================
    # FOOTER
    # ============================================================
    st.markdown("""
    <div class="footer">
        <div class="footer-content">
            <div class="footer-author">Developed by Dr. Volkan OBAN</div>
            <div class="footer-year">2025 - All Rights Reserved</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
if __name__ == "__main__":
    main()